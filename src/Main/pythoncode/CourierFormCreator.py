
from lib2to3.pgen2.token import SLASH
from docx import Document
import os

# get working directory
import sys

for arg in sys.argv:
    if arg.startswith("--dir="):
        working_directory = arg.split("=")[1]
        break
# raise Exception(working_directory)
working_directory.replace('\\', '/')
# working_directory = os.getcwd()
# replace SLASH with BACKSLASH
# working_directory = "C:/Git-hub/Post-Maroc/"#working_directory.replace('\\', '/')
print(working_directory)
class DocxCreator ():
    variables = {
        "${RRXXXXXXXXXMA}": "No information",
        "${LN1}": "No information",
        "${FN1}": "No information",
        "${LN2}": "No information",
        "${FN2}": "No information",
        "${ADDRESS1}": "No information",
        "${ADDRESS2}": "No information",
        "${PHONE1}": "No information",
        "${PHONE2}": "No information",
        }
    template_file_path = working_directory+'\\CourierForm.docx'
    output_file_path = working_directory+'\\result.docx'
    # read information from file txt
    def read_file(self):
        # print working directory
        print(os.getcwd())
        
        with open(working_directory+'\\CurrentCourierInfo.txt', 'r') as file:
            lines = file.readlines()
            # remove \n
            lines = [line.strip() for line in lines]
            return lines
    # constructor
    def __init__(self):
        readed_info = self.read_file()
        # stroe information in variables
        for i in range(len(readed_info)):
            self.variables[list(self.variables.keys())[i]] = readed_info[i]
        self.template_document = Document(self.template_file_path)
    def FillDocx(self):
        print(dir(self.template_document))
        for variable_key, variable_value in self.variables.items():
            for paragraph in self.template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
            for table in self.template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
    def save(self):
        self.template_document.add_picture(working_directory+"\\Barcode.png")
        self.template_document.save(self.output_file_path)

    def replace_text_in_paragraph(self,paragraph, key, value):
        # print(dir(paragraph))
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)
docx_creator = DocxCreator()
docx_creator.FillDocx()
docx_creator.save()